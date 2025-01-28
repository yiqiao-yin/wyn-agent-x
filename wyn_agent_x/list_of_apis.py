from typing import Any, Dict, List

from serpapi import GoogleSearch
from twilio.rest import Client

from wyn_agent_x.helper import register_function


@register_function("send_sms")
def send_sms(
    payload: Dict[str, str], secrets: Dict[str, str], event_stream: list
) -> Dict[str, Any]:
    """
    Simulate sending an SMS using the Twilio API.

    Args:
        payload (Dict[str, str]): Contains the message details and recipient information.
        secrets (Dict[str, str]): Contains sensitive data like account SID and auth token.
        event_stream (list): A list to log events and responses for the SMS API call.

    Returns:
        Dict[str, Any]: A dictionary containing the status and result of the SMS operation.

    The function limits the message body to 1600 characters.
    If exceeded, it returns an error response without sending the SMS.
    """
    print(f"👀 API Call: Sending SMS with payload: {payload}")

    # Extract Secrets
    account_sid = secrets["account_sid"]
    auth_token = secrets["auth_token"]

    # Message body check for character limit
    message_body = (
        f"Hello {payload['name']}, here's the message: {payload['message body']}"
    )
    if len(message_body) > 1600:
        response = {
            "status": "error",
            "message": "Message body exceeds 1600 character limit.",
            "model_name": "None",
        }
        print("👀 Error: Message body exceeds the 1600 character limit.")

        # Append the error result to the event stream
        event_stream.append(
            {"event": "api_call", "api_name": "send_sms", "response": response}
        )
        return response

    try:
        # Initialize Twilio Client with provided credentials
        client = Client(account_sid, auth_token)

        # Simulate sending a message (replace with actual logic for real SMS sending)
        message = client.messages.create(
            body=message_body,
            from_="+18552060350",  # Replace with a valid Twilio number
            to="+15859538396",  # Replace with the destination number
        )

        print(f"👀 Message SID: {message.sid}")
        response = {"status": f"success: {message.sid}", "model_name": "None"}

    except Exception as e:
        # Gracefully handle any exceptions
        response = {
            "status": f"error: {str(e)}",
            "message": "Failed to send SMS due to an error.",
            "model_name": "None",
        }
        print(f"👀 Error: {str(e)}")

    # Append the result to the event stream
    event_stream.append(
        {"event": "api_call", "api_name": "send_sms", "response": response}
    )

    return response


# Register the google_search function
@register_function("google_search")
def google_search(
    payload: Dict[str, str], secrets: Dict[str, str], event_stream: list
) -> Dict[str, Any]:
    """
    Simulate a Google search using the SerpAPI and return the results formatted as a Markdown table.

    Args:
        payload (Dict[str, str]): A dictionary containing the search parameters:
                                  - query: The search query string.
                                  - location: The location for the search.
        secrets (Dict[str, str]): Contains the SerpAPI key for API authentication.
        event_stream (list): A list used to log events and responses.

    Returns:
        Dict[str, Any]: A dictionary with the status of the API call and a Markdown table of search results.
    """

    # Extract Secrets
    serpapi_key = secrets["serpapi_key"]

    # Extract search parameters from the payload
    engine = payload.get("engine", "google")
    query = payload.get("query", "")
    location = payload.get("location", "")
    num = payload.get("num", "50")
    num = int(num)
    num = num if type(num) == int else 50

    # Checkpoint
    print(f"Run google search API using: \nquery={query}, \nlocation={location}")

    try:
        # Perform the Google Search using SerpAPI
        search = GoogleSearch(
            {
                "engine": engine,
                "q": query,
                "location": location,
                "api_key": serpapi_key,
                "num": num,
            }
        )
        results = search.get_dict()

        # Extract relevant results
        organic_results = results.get("organic_results", [])

        # Convert search results to a Markdown table
        md_table = "| Title | Link | Snippet |\n| :--- | :--- | :--- |\n"
        for item in organic_results:
            title = item.get("title", "N/A")
            link = item.get("link", "#")
            snippet = item.get("snippet", "N/A")
            md_table += f"| {title} | [Link]({link}) | {snippet} |\n"
        print(f"👀 Results: \n{md_table}")

        # Prepare the response
        response = {"status": "success", "model_name": "None", "results": md_table}

        # Append the result to the event stream
        event_stream.append(
            {
                "event": "api_call",
                "api_name": "google_search",
                "response": {"text": response, "status": "200 success"},
            }
        )

    except Exception as e:
        # Handle any exceptions that occur during the API call
        response = {"status": f"error: {str(e)}", "model_name": "None"}
        event_stream.append(
            {"event": "api_call", "api_name": "google_search", "response": response}
        )

    return response


import json  # To format dictionary as text
import os
import smtplib
from email.message import EmailMessage

import matplotlib.pyplot as plt
import pandas as pd  # To handle DataFrame type
import yfinance as yf
from fpdf import FPDF


@register_function("generate_financial_report")
def generate_financial_report(
    payload: Dict[str, str], secrets: Dict[str, str], event_stream: List[Dict[str, Any]]
) -> Dict[str, Any]:
    """
    Generate a financial report for a given ticker and save it as a PDF.

    Args:
        payload (Dict[str, str]): Contains the ticker symbol and save_pdf flag:
                                  - ticker: The stock ticker symbol (e.g., "AAPL").
                                  - save_pdf: Whether to save the report as a PDF (True/False).
        secrets (Dict[str, str]): Not used but included for consistency.
        event_stream (list): A list to log events and responses.

    Returns:
        Dict[str, Any]: A dictionary containing the status and the filename if saved.
    """
    # Extract payload values
    ticker = payload.get("ticker", "AAPL")
    save_pdf = payload.get("save_pdf", True)
    send_email = payload.get("send_email", True)
    to_email = payload.get("to_email", "eagle0504@gmail.com")

    # Extract secrets
    email_key = secrets["email_key"]

    # Log API call details
    print(f"Generating financial report for ticker: {ticker}, save_pdf: {save_pdf}")

    # Fetching data from Yahoo Finance
    result = yf.Ticker(ticker)
    historical_data = result.history(period="10y")
    analyst_price_targets = result.analyst_price_targets
    recommendations_summary = result.recommendations_summary
    info = result.info

    # Extract relevant info for the report
    long_business_summary = info.get("longBusinessSummary", "No summary available.")
    sector = info.get("sector", "No sector information available.")
    website = info.get("website", "No website information available.")

    # Function to format data as text
    def format_data(data):
        if data is None:
            return "No data available."
        elif isinstance(data, pd.DataFrame):
            return data.to_string(index=False)
        elif isinstance(data, dict):
            return json.dumps(data, indent=4)
        else:
            return str(data)

    # Function to calculate RSI
    def calculate_rsi(data, window=14):
        delta = data.diff()
        gain = (delta.where(delta > 0, 0)).rolling(window=window).mean()
        loss = (-delta.where(delta < 0, 0)).rolling(window=window).mean()
        rs = gain / loss
        return 100 - (100 / (1 + rs))

    # Function to plot the stock chart
    def plot_stock_chart(data, ticker, filename, analyst_price_targets):
        plt.figure(figsize=(10, 8))

        # Subplot 1: Stock price with moving averages
        plt.subplot(2, 1, 1)
        plt.plot(data.index, data["Close"], label="Close Price", color="blue")
        plt.plot(data.index, data["SMA20"], label="SMA20", color="orange")
        plt.plot(data.index, data["SMA50"], label="SMA50", color="green")
        plt.plot(data.index, data["SMA100"], label="SMA100", color="red")

        # Add analyst price targets as stars
        if analyst_price_targets:
            last_date = data.index[-1]
            next_date = last_date + pd.Timedelta(
                days=1
            )  # Assume the next timestamp is the next day
            plt.scatter(
                next_date,
                analyst_price_targets["current"],
                color="gold",
                label="Current Price Target",
                s=200,
                marker="*",
            )
            plt.scatter(
                next_date,
                analyst_price_targets["low"],
                color="red",
                label="Low Price Target",
                s=200,
                marker="*",
            )
            plt.scatter(
                next_date,
                analyst_price_targets["high"],
                color="green",
                label="High Price Target",
                s=200,
                marker="*",
            )
            plt.scatter(
                next_date,
                analyst_price_targets["mean"],
                color="blue",
                label="Mean Price Target",
                s=200,
                marker="*",
            )
            plt.scatter(
                next_date,
                analyst_price_targets["median"],
                color="purple",
                label="Median Price Target",
                s=200,
                marker="*",
            )

        plt.title(f"{ticker} Stock Price and Moving Averages")
        plt.legend()
        plt.grid()

        # Subplot 2: RSI
        plt.subplot(2, 1, 2)
        plt.plot(data.index, data["RSI"], label="RSI", color="purple")
        plt.axhline(70, color="red", linestyle="--", linewidth=1)
        plt.axhline(30, color="green", linestyle="--", linewidth=1)
        plt.title("Relative Strength Index (RSI)")
        plt.legend()
        plt.grid()

        plt.tight_layout()
        plt.savefig(filename, format="png")  # Save the chart as a PNG file
        plt.close()

    # Add moving averages and RSI to the historical data
    historical_data["SMA20"] = historical_data["Close"].rolling(window=20).mean()
    historical_data["SMA50"] = historical_data["Close"].rolling(window=50).mean()
    historical_data["SMA100"] = historical_data["Close"].rolling(window=100).mean()
    historical_data["RSI"] = calculate_rsi(historical_data["Close"])

    response = {}

    try:
        if save_pdf:
            # Create PDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", "B", 16)
            pdf.cell(200, 10, f"Financial Summary for {ticker}", ln=True, align="C")

            # Add business overview
            pdf.ln(10)
            pdf.set_font("Arial", "B", 12)
            pdf.cell(200, 10, "Business Overview", ln=True)
            pdf.set_font("Arial", "", 10)
            pdf.multi_cell(0, 10, f"Sector: {sector}")
            pdf.multi_cell(0, 10, f"Website: {website}")
            pdf.multi_cell(0, 10, f"Summary: {long_business_summary}")

            # Plot stock chart and save it as a PNG
            chart_filename = f"{ticker}_stock_chart.png"
            plot_stock_chart(
                historical_data, ticker, chart_filename, analyst_price_targets
            )

            # Add the stock chart to the PDF
            pdf.add_page()
            pdf.set_font("Arial", "B", 12)
            pdf.cell(200, 10, "Stock Chart", ln=True)
            pdf.image(chart_filename, x=10, y=30, w=180)  # Add the PNG image

            # Save PDF
            pdf_filename = f"{ticker}_report.pdf"
            pdf.output(pdf_filename)
            print(f"PDF saved as {pdf_filename}.")
            response = {"status": "success", "file": pdf_filename}

            if send_email:
                # Param required for sending email
                file_name = pdf_filename
                # to_email = to_email

                # Check if the file is a PDF
                if not file_name.endswith(".pdf"):
                    raise ValueError("The file must be a .pdf file")

                # Check if the file exists
                if not os.path.exists(file_name):
                    raise FileNotFoundError(
                        f"File '{file_name}' not found in the local directory"
                    )

                # Email credentials (use environment variables or replace directly)
                from_email = "eagle0504@gmail.com"  # Your email address
                password = email_key  # Your email password (use app-specific password for security)

                # Create email message
                msg = EmailMessage()
                msg["Subject"] = "Sample PDF Report"
                msg["From"] = from_email
                msg["To"] = to_email
                msg.set_content("Please find the attached PDF file.")

                # Attach the PDF file
                with open(file_name, "rb") as f:
                    pdf_data = f.read()
                    msg.add_attachment(
                        pdf_data,
                        maintype="application",
                        subtype="pdf",
                        filename=file_name,
                    )

                # Send the email
                try:
                    with smtplib.SMTP_SSL(
                        "smtp.gmail.com", 465
                    ) as server:  # SMTP server for Gmail
                        server.login(from_email, password)
                        server.send_message(msg)
                    print(
                        f"Email sent successfully to {to_email} with the attached PDF: {file_name}"
                    )
                except Exception as e:
                    print(f"Failed to send email: {str(e)}")
        else:
            response = {"status": "skipped", "message": "PDF saving is disabled."}

    except Exception as e:
        response = {
            "status": f"error: {str(e)}",
            "message": "Failed to generate financial report.",
        }
        print(f"Error: {str(e)}")

    # Append result to event stream
    event_stream.append(
        {
            "event": "api_call",
            "api_name": "generate_financial_report",
            "response": response,
        }
    )

    return response


from docx import Document


@register_function("generate_word_document")
def generate_word_document(
    payload: Dict[str, str], secrets: Dict[str, str], event_stream: List[Dict[str, Any]]
) -> Dict[str, Any]:
    """
    Generate a Word document with a given header and paragraph and optionally send it via email.

    Args:
        payload (Dict[str, str]): Contains the header, paragraph, and email-related flags:
                                  - header: The title for the Word document.
                                  - paragraph: The content of the document.
                                  - save_pdf: Whether to save the Word document as a .docx file (True/False).
                                  - send_email: Whether to send the document via email (True/False).
                                  - to_email: Recipient email address.
        secrets (Dict[str, str]): Contains the email credentials (e.g., email_key).
        event_stream (list): A list to log events and responses.

    Returns:
        Dict[str, Any]: A dictionary containing the status and the filename if saved.
    """
    # Extract payload values
    header = payload.get("header", "Default Title")
    paragraph = payload.get("paragraph", "Default paragraph content.")
    save_docx = payload.get("save_docx", True)
    send_email = payload.get("send_email", True)
    to_email = payload.get("to_email", "eagle0504@gmail.com")

    # Extract secrets
    email_key = secrets.get("email_key")

    # Log the API call details
    print(f"Generating Word document with header: {header}, save_docx: {save_docx}")

    response = {}

    try:
        if save_docx:
            # Create a Word document
            document = Document()
            document.add_heading(header, level=1)
            document.add_paragraph(paragraph)

            # Save the document
            file_name = "generated_document.docx"
            document.save(file_name)
            print(f"Word document saved as {file_name}.")
            response = {"status": "success", "file": file_name}

            if send_email:
                # Check if the file exists
                if not os.path.exists(file_name):
                    raise FileNotFoundError(
                        f"File '{file_name}' not found in the local directory"
                    )

                # Email credentials (use environment variables or replace directly)
                from_email = "eagle0504@gmail.com"  # Your email address
                password = email_key  # Your email password (use app-specific password for security)

                # Create email message
                msg = EmailMessage()
                msg["Subject"] = "Generated Word Document"
                msg["From"] = from_email
                msg["To"] = to_email
                msg.set_content("Please find the attached Word document.")

                # Attach the Word document
                with open(file_name, "rb") as f:
                    file_data = f.read()
                    msg.add_attachment(
                        file_data,
                        maintype="application",
                        subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
                        filename=file_name,
                    )

                # Send the email
                try:
                    with smtplib.SMTP_SSL(
                        "smtp.gmail.com", 465
                    ) as server:  # SMTP server for Gmail
                        server.login(from_email, password)
                        server.send_message(msg)
                    print(
                        f"Email sent successfully to {to_email} with the attached document: {file_name}"
                    )
                except Exception as e:
                    print(f"Failed to send email: {str(e)}")

        else:
            response = {
                "status": "skipped",
                "message": "Saving the Word document is disabled.",
            }

    except Exception as e:
        response = {
            "status": f"error: {str(e)}",
            "message": "Failed to generate Word document.",
        }
        print(f"Error: {str(e)}")

    # Append result to event stream
    event_stream.append(
        {
            "event": "api_call",
            "api_name": "generate_word_document",
            "response": response,
        }
    )

    return response


import base64
import os

import anthropic
import pyautogui


@register_function("read_screen")
def read_screen(
    payload: Dict[str, Any], secrets: Dict[str, str], event_stream: List[Dict[str, Any]]
) -> Dict[str, Any]:
    """
    Takes a screenshot if requested, then optionally calls the Claude API with the screenshot
    (converted to Base64) for an image caption. Returns the results in a dictionary.

    Args:
        payload (Dict[str, Any]): A dictionary containing:
            - read_screen (bool): If True, take a screenshot using pyautogui.
            - call_claude_api (bool): If True, send the screenshot to the Claude API for a caption.
        secrets (Dict[str, str]): A dictionary containing:
            - CLAUDE_API_KEY: The API key for Claude (retrieved from secrets).
        event_stream (List[Dict[str, Any]]): A list to log events and responses.

    Returns:
        Dict[str, Any]: A dictionary containing the operation status and any captured data, for example:
            {
                "screenshot_taken": True/False,
                "img_caption": "String describing the screenshot" or "N/A"
            }
    """

    # Extract payload flags
    read_screen_flag: bool = payload.get("read_screen", False)
    call_claude_api_flag: bool = payload.get("call_claude_api", False)

    # Retrieve Claude API key from secrets
    claude_api_key: str = secrets.get("claude_api_key", "")

    response: Dict[str, Any] = {}

    # Take screenshot if requested
    if read_screen_flag:
        screenshot_file: str = "my_screenshot.png"
        # Capture screenshot
        im1 = pyautogui.screenshot()
        im1.save(screenshot_file)
        print(f"Screenshot saved as {screenshot_file}.")
        response["screenshot_taken"] = True
    else:
        screenshot_file: str = ""
        response["screenshot_taken"] = False

    # Call Claude API if requested and we have an API key
    if call_claude_api_flag and claude_api_key:
        try:
            # Encode screenshot to Base64 only if we took one
            if response["screenshot_taken"]:
                with open(screenshot_file, "rb") as f:
                    file_obj: str = base64.standard_b64encode(f.read()).decode("utf-8")
            else:
                # If no screenshot was taken, you can choose to handle differently (e.g., skip).
                file_obj: str = ""

            # Initialize the Claude client
            client = anthropic.Anthropic(api_key=claude_api_key)

            # Make a request to Claude, embedding the screenshot as Base64 data
            message = client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=1024,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image",
                                "source": {
                                    "type": "base64",
                                    "media_type": "image/png",
                                    "data": file_obj,
                                },
                            },
                            {"type": "text", "text": "Describe this image."},
                        ],
                    }
                ],
            )

            # Extract the image caption from Claude's response
            img_caption: str = message.content[0].text
            response["img_caption"] = img_caption
            print(f"Claude API caption: {img_caption}")

        except Exception as e:
            error_msg = f"Error calling Claude API: {str(e)}"
            response["img_caption"] = error_msg
            print(error_msg)
    else:
        response["img_caption"] = "N/A"
        if call_claude_api_flag and not claude_api_key:
            print("Claude API call requested, but no API key was provided.")

    # Log the operation to the event stream
    event_stream.append(
        {
            "event": "api_call",
            "api_name": "read_screen",
            "response": response,
        }
    )

    return response


import pyautogui
import requests
import json
import base64
from typing import Dict, Any, List

@register_function("find_and_click")
def find_and_click(
    payload: Dict[str, Any],
    secrets: Dict[str, str],
    event_stream: List[Dict[str, Any]]
) -> Dict[str, Any]:
    """
    Orchestrates a full workflow to:
      1) Capture a screenshot of the entire screen,
      2) Convert it to Base64,
      3) POST it to an AWS Textract-like endpoint to get OCR data,
      4) Clean the OCR data,
      5) Find the OCR entry whose text best matches the 'target_location' string from the payload,
      6) Convert that entry's bounding box (in percentage) to an absolute screen region,
      7) Screenshot that region and save it,
      8) Locate that smaller screenshot on the full screen,
      9) Move the mouse there and click.

    Args:
        payload (Dict[str, Any]): A dictionary expected to contain:
            - target_location (str): The text we want to find in the OCR results.
        secrets (Dict[str, str]): A dictionary of secrets (not used in this function, included for parity).
        event_stream (List[Dict[str, Any]]): A list to log events and responses.

    Returns:
        Dict[str, Any]: A dictionary containing:
            {
                "screenshot_taken": bool,
                "screenshot_saved": bool,
                "screenshot_file": str,
                "image_found_on_screen": bool,
                "click_performed": bool,
                "error": str (optional),
            }
    """

    # Initialize a response dictionary
    response: Dict[str, Any] = {
        "screenshot_taken": False,
        "screenshot_saved": False,
        "screenshot_file": "",
        "image_found_on_screen": False,
        "click_performed": False
    }

    # Extract the single required value from payload
    target_location: str = payload.get("target_location", "")

    # Validate minimal input
    if not target_location:
        error_msg = (
            "Missing required payload parameter. 'target_location' must be provided."
        )
        print(error_msg)
        response["error"] = error_msg
        event_stream.append(
            {
                "event": "api_call",
                "api_name": "find_and_click",
                "response": response,
            }
        )
        return response

    # Helper functions defined locally:

    def image_to_base64(file_path: str) -> str:
        """
        Reads an image file from the given file path and returns its content encoded in base64.

        Args:
            file_path (str): The path to the image file (e.g., png, jpg).

        Returns:
            str: The base64 encoded content of the image file, or an empty string if an error occurs.
        """
        try:
            with open(file_path, "rb") as image_file:
                image_content = image_file.read()
            base64_encoded_image = base64.b64encode(image_content)
            return base64_encoded_image.decode("utf-8")
        except Exception as e:
            print(f"Error in image_to_base64: {str(e)}")
            return ""

    def post_request_and_parse_response(url: str, data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Sends a POST request to the specified URL with the given data,
        then parses the byte response to a dictionary.

        Args:
            url (str): The URL to which the POST request is sent.
            data (Dict[str, Any]): The data to send in the POST request.

        Returns:
            Dict[str, Any]: The parsed dictionary from the response or an empty dict on error.
        """
        headers = {"Content-Type": "application/json"}
        try:
            response = requests.post(url, json=data, headers=headers)
            byte_data = response.content
            decoded_string = byte_data.decode("utf-8")
            dict_data = json.loads(decoded_string)
            return dict_data
        except Exception as ex:
            print(f"Error in post_request_and_parse_response: {str(ex)}")
            return {}

    def clean_ocr_results(ocr_data: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Cleans the OCR results obtained from AWS Textract by extracting the text
        and bounding box for each line that contains a 'Text' field.

        Args:
            ocr_data (List[Dict[str, Any]]): A list of OCR result dictionaries as parsed from JSON.

        Returns:
            List[Dict[str, Any]]: A list of dictionaries, each containing 'text' and 'bbox'
                                  (bounding box) for valid OCR entries.
        """
        cleaned_list = []
        for item in ocr_data:
            if "Text" in item:
                bbox = item.get("Geometry", {}).get("BoundingBox", {})
                cleaned_entry = {
                    "text": item["Text"],
                    "bbox": bbox
                }
                cleaned_list.append(cleaned_entry)
        return cleaned_list

    def find_closest_match(ocr_results: List[Dict[str, Any]], prompt: str) -> Dict[str, Any]:
        """
        Finds the OCR result whose 'text' is the closest match to a given prompt,
        using the Levenshtein distance metric.

        Args:
            ocr_results (List[Dict[str, Any]]): A list of dictionaries, each containing at least:
                {
                  "text": <string>,
                  "bbox": {
                    "Width": <float>,
                    "Height": <float>,
                    "Left": <float>,
                    "Top": <float>
                  }
                }
            prompt (str): The string to match against the OCR text.

        Returns:
            Dict[str, Any]: The dictionary from 'ocr_results' whose 'text' field has
                            the smallest Levenshtein distance to the prompt,
                            or an empty dict if no results exist.
        """
        def levenshtein_distance(str1: str, str2: str) -> int:
            if not str1:
                return len(str2)
            if not str2:
                return len(str1)

            dp = [[0] * (len(str2) + 1) for _ in range(len(str1) + 1)]
            for i in range(len(str1) + 1):
                dp[i][0] = i
            for j in range(len(str2) + 1):
                dp[0][j] = j

            for i in range(1, len(str1) + 1):
                for j in range(1, len(str2) + 1):
                    if str1[i - 1] == str2[j - 1]:
                        dp[i][j] = dp[i - 1][j - 1]
                    else:
                        dp[i][j] = 1 + min(
                            dp[i - 1][j],    # Deletion
                            dp[i][j - 1],    # Insertion
                            dp[i - 1][j - 1] # Substitution
                        )
            return dp[len(str1)][len(str2)]

        min_distance = float("inf")
        best_match = {}
        for item in ocr_results:
            candidate_text = item.get("text", "")
            distance = levenshtein_distance(candidate_text.lower(), prompt.lower())
            if distance < min_distance:
                min_distance = distance
                best_match = item
        return best_match

    try:
        # 1) Capture a screenshot of the entire screen
        screen_width, screen_height = pyautogui.size()

        # Calculate 75% of the screen width
        width_75pct = int(screen_width * 0.75)

        # The region is defined as (left, top, width, height).
        region = (0, 0, width_75pct, screen_height)

        # Capture a screenshot of the specified region
        partial_screenshot = pyautogui.screenshot(region=region)
        screenshot_file: str = "whole_screenshot.png"
        im1 = partial_screenshot
        im1.save(screenshot_file)
        print(f"Screenshot of whole screen saved as {screenshot_file}.")
        response["screenshot_taken"] = True

        # 2) Convert this screenshot to Base64
        image_base64: str = image_to_base64(screenshot_file)
        if not image_base64:
            error_msg = "Could not convert the screenshot to Base64. Aborting."
            print(error_msg)
            response["error"] = error_msg
            event_stream.append(
                {
                    "event": "api_call",
                    "api_name": "find_and_click",
                    "response": response,
                }
            )
            return response

        # 3) POST this base64 image to the Textract-like endpoint
        url = "https://2tsig211e0.execute-api.us-east-1.amazonaws.com/my_textract"
        post_data = {"image": image_base64}
        result_dict = post_request_and_parse_response(url, post_data)
        print("OCR response received. Here are the keys:")
        print(result_dict.keys())

        # 4) Check if 'body' is in the response dictionary
        if "body" not in result_dict:
            error_msg = "No 'body' found in the OCR response. Aborting."
            print(error_msg)
            response["error"] = error_msg
            event_stream.append(
                {
                    "event": "api_call",
                    "api_name": "find_and_click",
                    "response": response,
                }
            )
            return response

        # 5) Parse the 'body' field, which should be a JSON string
        ocr_data: List[Dict[str, Any]] = json.loads(result_dict["body"])

        # 6) Clean up the raw OCR results
        cleaned_up_results: List[Dict[str, Any]] = clean_ocr_results(ocr_data)

        # 7) Find the best match for the text (using 'target_location' as prompt)
        best_matched_result: Dict[str, Any] = find_closest_match(cleaned_up_results, target_location)
        if not best_matched_result:
            error_msg = "No suitable OCR match found. Aborting."
            print(error_msg)
            response["error"] = error_msg
            event_stream.append(
                {
                    "event": "api_call",
                    "api_name": "find_and_click",
                    "response": response,
                }
            )
            return response

        # 8) Extract bounding box percentages from the best match
        bbox: Dict[str, float] = best_matched_result.get("bbox", {})
        width_pct: float = bbox.get("Width", 0.0)
        height_pct: float = bbox.get("Height", 0.0)
        left_pct: float = bbox.get("Left", 0.0)
        top_pct: float = bbox.get("Top", 0.0)

        # 9) Convert percentages to absolute screen coordinates
        screen_width, screen_height = pyautogui.size()
        abs_left = int(left_pct * screen_width * 0.75)  # 75% of the screen width
        abs_top = int(top_pct * screen_height)
        abs_width = int(width_pct * screen_width * 0.75)  # 75% of the screen width
        abs_height = int(height_pct * screen_height)

        region = (abs_left, abs_top, abs_width, abs_height)

        # 10) Take a screenshot of the specified region
        screenshot_path = "screenshot_region.png"
        screenshot = pyautogui.screenshot(region=region)
        screenshot.save(screenshot_path)
        print(f"Screenshot of the specified region saved as '{screenshot_path}'.")
        response["screenshot_saved"] = True
        response["screenshot_file"] = screenshot_path

        # 11) Locate the center of the saved screenshot on the screen
        center = pyautogui.locateCenterOnScreen(screenshot_path)
        if center:
            x, y = center
            print(f"Screenshot image found at coordinates: ({x}, {y}).")
            response["image_found_on_screen"] = True

            # Move mouse to the found location (3-second duration)
            pyautogui.moveTo(x, y, duration=3)
            print("Mouse moved to the center of the screenshot image over 3 seconds.")

            # Click on the found center
            pyautogui.click()
            print("Clicked on the screenshot image.")
            response["click_performed"] = True
        else:
            print(f"Screenshot image ('{screenshot_path}') not found on the screen.")

    except Exception as e:
        error_msg = f"Error in find_and_click: {str(e)}"
        print(error_msg)
        response["error"] = error_msg

    # Log the operation to the event_stream
    event_stream.append(
        {
            "event": "api_call",
            "api_name": "find_and_click",
            "response": {"text": response, "status": "200 success"},
        }
    )

    return response


@register_function("open_google_browser")
def open_google_browser(
    payload: Dict[str, Any],
    secrets: Dict[str, str],
    event_stream: list
) -> Dict[str, Any]:
    """
    Opens the specified URL in the default web browser if the user confirms.

    Args:
        payload (Dict[str, Any]): A dictionary containing the user parameters:
                                  - verify_open_browser (bool): Whether to open the browser or not.
                                  - target_url_to_open (str): The URL to be opened in the browser.
        secrets (Dict[str, str]): Not used for this function, but present to maintain interface consistency.
        event_stream (list): A list used to log events and responses.

    Returns:
        Dict[str, Any]: A dictionary with the status of the operation and a message.
    """
    import webbrowser

    def open_url(url: str) -> None:
        """
        Opens the specified URL in the user's default web browser.

        Args:
            url (str): The URL to open.
        """
        # The webbrowser.open() function attempts to open the provided URL
        # in the default browser on the user's system
        webbrowser.open(url)

    # Extract the required payload parameters
    target_url_to_open = payload.get("target_url_to_open", "https://www.google.com")
    verify_open_browser = payload.get("verify_open_browser", False)

    if verify_open_browser:
        # Open the user-specified URL
        open_url(target_url_to_open)
        response = {
            "status": "success",
            "model_name": "None",
            "message": f"Browser opened to {target_url_to_open}.",
        }
    else:
        # No action taken
        response = {
            "status": "no_action",
            "model_name": "None",
            "message": "User declined to open the browser.",
        }

    # Log the event and response
    event_stream.append(
        {
            "event": "api_call",
            "api_name": "open_google_browser",
            "response": {"text": response, "status": "200 success"},
        }
    )

    return response


from typing import Dict, Any, List
import pyautogui

@register_function("enter_text_on_screen")
def enter_text_on_screen(
    payload: Dict[str, str],
    secrets: Dict[str, str],
    event_stream: List[Any]
) -> Dict[str, Any]:
    """
    Simulate typing text on the screen using the PyAutoGUI library.

    Args:
        payload (Dict[str, str]): A dictionary containing:
                                  - text_to_write: The text to write on the screen.
        secrets (Dict[str, str]): Reserved for any secret values or API keys (not used here).
        event_stream (List[Any]): A list used to log events and responses.

    Returns:
        Dict[str, Any]: A dictionary with the status of the call and any relevant response data.
    """
    text_to_write = payload.get("text_to_write", "")

    try:
        # Click on the current location to ensure the text is entered in the right place
        pyautogui.click()

        # Use PyAutoGUI to write the specified text
        pyautogui.write(text_to_write)

        # Prepare the success response
        response = {
            "status": "success",
            "model_name": "None",
            "message": f"Typed text: {text_to_write}"
        }

        # Log event in the event stream
        event_stream.append(
            {
                "event": "api_call",
                "api_name": "enter_text_on_screen",
                "response": {"text": response, "status": "200 success"},
            }
        )
    except Exception as e:
        # Handle exceptions gracefully
        response = {"status": f"error: {str(e)}", "model_name": "None"}
        event_stream.append(
            {
                "event": "api_call",
                "api_name": "enter_text_on_screen",
                "response": response
            }
        )

    return response


from typing import Dict, Any, List
import pyautogui

@register_function("scroll_up")
def scroll_up(
    payload: Dict[str, str],
    secrets: Dict[str, str],
    event_stream: List[Any]
) -> Dict[str, Any]:
    """
    Scroll up by a specified number of lines using PyAutoGUI.

    Args:
        payload (Dict[str, str]): A dictionary containing:
                                  - number_of_lines_to_scroll: (optional) The number of lines to scroll up.
        secrets (Dict[str, str]): Reserved for any secret values or API keys (not used here).
        event_stream (List[Any]): A list used to log events and responses.

    Returns:
        Dict[str, Any]: A dictionary with the status of the call and any relevant response data.
    """
    # Extract number_of_lines_to_scroll from the payload, default to 7 if not provided or invalid
    raw_scroll_value = payload.get("number_of_lines_to_scroll", 7)
    try:
        number_of_lines_to_scroll = int(raw_scroll_value)
    except (ValueError, TypeError):
        number_of_lines_to_scroll = 7

    try:
        # Perform the scroll action using PyAutoGUI
        pyautogui.scroll(number_of_lines_to_scroll)

        # Prepare the success response
        response = {
            "status": "success",
            "model_name": "None",
            "message": f"Scrolled up {number_of_lines_to_scroll} lines."
        }

        # Log event in the event stream
        event_stream.append(
            {
                "event": "api_call",
                "api_name": "scroll_up",
                "response": {"text": response, "status": "200 success"},
            }
        )
    except Exception as e:
        # Handle exceptions gracefully
        response = {"status": f"error: {str(e)}", "model_name": "None"}
        event_stream.append(
            {
                "event": "api_call",
                "api_name": "scroll_up",
                "response": response
            }
        )

    return response


@register_function("scroll_down")
def scroll_down(
    payload: Dict[str, str],
    secrets: Dict[str, str],
    event_stream: List[Any]
) -> Dict[str, Any]:
    """
    Scroll down by a specified number of lines using PyAutoGUI.

    Args:
        payload (Dict[str, str]): A dictionary containing:
                                  - number_of_lines_to_scroll: (optional) The number of lines to scroll down.
        secrets (Dict[str, str]): Reserved for any secret values or API keys (not used here).
        event_stream (List[Any]): A list used to log events and responses.

    Returns:
        Dict[str, Any]: A dictionary with the status of the call and any relevant response data.
    """
    # Extract number_of_lines_to_scroll from the payload, default to 7 if not provided or invalid
    raw_scroll_value = payload.get("number_of_lines_to_scroll", 7)
    try:
        number_of_lines_to_scroll = int(raw_scroll_value)
    except (ValueError, TypeError):
        number_of_lines_to_scroll = 7

    try:
        # Perform the scroll action (negative value for scrolling down)
        pyautogui.scroll(-number_of_lines_to_scroll)

        # Prepare the success response
        response = {
            "status": "success",
            "model_name": "None",
            "message": f"Scrolled down {number_of_lines_to_scroll} lines."
        }

        # Log event in the event stream
        event_stream.append(
            {
                "event": "api_call",
                "api_name": "scroll_down",
                "response": {"text": response, "status": "200 success"},
            }
        )
    except Exception as e:
        # Handle exceptions gracefully
        response = {"status": f"error: {str(e)}", "model_name": "None"}
        event_stream.append(
            {
                "event": "api_call",
                "api_name": "scroll_down",
                "response": response
            }
        )

    return response


from typing import Dict, Any, List
import pyautogui

@register_function("click_now")
def click_now(
    payload: Dict[str, str],
    secrets: Dict[str, str],
    event_stream: List[Any]
) -> Dict[str, Any]:
    """
    Perform a mouse click using the PyAutoGUI library.

    Args:
        payload (Dict[str, str]): Reserved for any arguments (not used in this function).
        secrets (Dict[str, str]): Reserved for any secret values or API keys (not used here).
        event_stream (List[Any]): A list used to log events and responses.

    Returns:
        Dict[str, Any]: A dictionary with the status of the call and any relevant response data.
    """
    try:
        # Perform a mouse click
        pyautogui.click()

        # Prepare the success response
        response = {
            "status": "success",
            "model_name": "None",
            "message": "Mouse click performed."
        }

        # Log event in the event stream
        event_stream.append(
            {
                "event": "api_call",
                "api_name": "click_now",
                "response": {"text": response, "status": "200 success"},
            }
        )
    except Exception as e:
        # Handle exceptions gracefully
        response = {"status": f"error: {str(e)}", "model_name": "None"}
        event_stream.append(
            {
                "event": "api_call",
                "api_name": "click_now",
                "response": response
            }
        )

    return response
